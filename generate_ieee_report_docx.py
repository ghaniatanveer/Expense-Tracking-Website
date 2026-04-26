from datetime import date
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            row_cells[idx].text = val


doc = Document()

# Title page
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.add_run("Expense Tracker Pro – A Smart Personal Finance Manager").bold = True
title.runs[0].font.size = None
doc.add_paragraph().add_run("IEEE Style Technical Report").bold = True
doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n")
doc.add_paragraph("Author: Your Name").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Course/Project: Final Year Project").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph(f"Date: {date.today().isoformat()}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

add_heading(doc, "Abstract", level=1)
add_para(
    doc,
    "Expense Tracker Pro is a full-stack web application for personal finance management. "
    "The platform combines a PHP REST backend and a Vanilla JavaScript frontend, with file-based JSON persistence. "
    "It supports expense/category CRUD, budgeting, recurring expense automation, report exports, interactive analytics, "
    "dark mode, and forecasting via linear regression. This report presents architecture, implementation, testing, "
    "limitations, and future work in IEEE technical style.",
)

add_heading(doc, "I. Introduction", level=1)
add_para(
    doc,
    "Manual expense tracking is often inconsistent and difficult to analyze over time. People lose visibility into where "
    "money is spent, whether budgets are respected, and how future spending may evolve. Expense Tracker Pro was developed "
    "to solve these problems with a practical and deployable web solution."
)
add_para(doc, "Project objectives:", bold=True)
add_bullets(doc, [
    "Implement robust CRUD for expenses and categories.",
    "Add budget monitoring with progress indicators and threshold alerts.",
    "Provide analytics and reporting with charts and exports.",
    "Automate recurring expenses.",
    "Integrate lightweight ML forecasting for next-month spending."
])

add_heading(doc, "II. System Architecture", level=1)
add_para(
    doc,
    "The architecture follows a two-layer model: (1) client-side presentation and logic layer, and (2) backend service layer. "
    "The frontend handles forms, filters, sorting, rendering, state synchronization, and chart generation. "
    "The backend exposes REST-like endpoints through backend.php and persists records in JSON files."
)
add_para(doc, "[Figure Placeholder: High-Level Architecture Diagram]")

add_heading(doc, "III. Features and Modules", level=1)
add_heading(doc, "A. Expense Management", level=2)
add_bullets(doc, [
    "Create, update, delete, and list expenses.",
    "Fields: name, amount, category, date, notes, tags, recurring flags.",
    "Search, sort, and date-range filtering in a tabular interface."
])
add_heading(doc, "B. Category and Budget Management", level=2)
add_bullets(doc, [
    "Category CRUD with color and icon metadata.",
    "Per-category monthly budget allocation.",
    "Overall monthly budget limit in settings."
])
add_heading(doc, "C. Reporting and Analytics", level=2)
add_bullets(doc, [
    "Pie chart for category distribution.",
    "Bar chart for six-month spending trend.",
    "Line chart for cumulative spending.",
    "CSV and PDF export support."
])
add_heading(doc, "D. Recurring Expense Module", level=2)
add_para(
    doc,
    "Recurring templates are stored with period and next_date. A recurring-run endpoint checks due items and automatically "
    "adds them into expenses, then advances next_date."
)
add_heading(doc, "E. ML Forecasting Module", level=2)
add_para(
    doc,
    "A linear regression model predicts next month spending from monthly aggregate history. "
    "A confidence indicator is derived from fit quality, and unusually large transactions are flagged as anomalies."
)

add_heading(doc, "IV. Data Design", level=1)
add_para(doc, "Data is stored in four JSON files:", bold=True)
add_bullets(doc, [
    "categories.json: id, name, color, budget_monthly, icon",
    "expenses.json: id, name, amount, category_id, date, notes, tags, is_recurring, recurring_period",
    "recurring.json: id, name, amount, category_id, period, next_date",
    "settings.json: currency, theme, default_view, monthly_budget_limit"
])
add_para(doc, "Sample JSON structures should be inserted here as code listings for final submission.")

add_heading(doc, "V. API Documentation", level=1)
add_table(
    doc,
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/backend.php?entity=expenses", "Fetch all expenses"],
        ["POST", "/backend.php?entity=expenses", "Create new expense"],
        ["PUT", "/backend.php?entity=expenses&id={id}", "Update expense"],
        ["DELETE", "/backend.php?entity=expenses&id={id}", "Delete expense"],
        ["GET", "/backend.php?entity=categories", "Fetch categories"],
        ["POST", "/backend.php?entity=categories", "Create category"],
        ["GET", "/backend.php?entity=recurring", "Fetch recurring rules"],
        ["POST", "/backend.php?entity=recurring-run", "Auto-generate due recurring expenses"],
        ["GET", "/backend.php?entity=settings", "Fetch settings"],
        ["PUT", "/backend.php?entity=settings", "Update settings"],
    ]
)
add_para(doc, "All responses follow a consistent envelope: success, message, and data.")

add_heading(doc, "VI. Machine Learning Implementation", level=1)
add_para(
    doc,
    "Feature engineering step aggregates expenses month-wise. For each month i, point (x=i, y=monthly_total) is generated. "
    "Least squares linear regression computes slope and intercept. Prediction is made for x=n+1. "
    "This approach is lightweight, interpretable, and suitable for client-side execution."
)
add_para(doc, "Limitations include linearity assumptions, low history quality, and absence of external economic factors.")

add_heading(doc, "VII. User Manual", level=1)
add_para(doc, "Installation and run steps:", bold=True)
add_bullets(doc, [
    "Set up a PHP-capable server environment.",
    "Place project files in web root.",
    "Ensure JSON files are writable.",
    "Run server and open index.php in browser.",
    "Use dashboard forms and table to manage expenses and categories."
])
add_para(doc, "[Figure Placeholder: Dashboard Screenshot]")

add_heading(doc, "VIII. Testing and Validation", level=1)
add_table(
    doc,
    ["Test ID", "Scenario", "Expected Result"],
    [
        ["T1", "Add expense", "Record created and visible in table"],
        ["T2", "Update expense", "Values update in UI and storage"],
        ["T3", "Delete expense", "Record removed from table and JSON"],
        ["T4", "Budget threshold", "Warning/exceeded visual alert shown"],
        ["T5", "Recurring run", "Due entries auto-added"],
        ["T6", "Forecast", "Prediction and confidence displayed"],
    ],
)
add_para(doc, "Manual validation confirmed functional behavior for major modules and edge cases.")

add_heading(doc, "IX. Challenges and Future Work", level=1)
add_bullets(doc, [
    "Challenge: file-based storage has limited concurrency guarantees.",
    "Challenge: single-file backend can become large without strict modularity.",
    "Future: migrate to SQLite/MySQL with transactions.",
    "Future: add authentication and multi-user profiles.",
    "Future: improve forecasting using ARIMA/LSTM models.",
    "Future: include automated CI tests and API contract tests."
])

add_heading(doc, "X. Conclusion", level=1)
add_para(
    doc,
    "Expense Tracker Pro demonstrates a professional full-stack implementation using a lightweight technology stack. "
    "It successfully combines operational expense tracking with analytical and predictive capabilities, providing both "
    "usability and technical extensibility for academic and real-world use."
)

add_heading(doc, "References", level=1)
refs = [
    "[1] PHP Documentation, https://www.php.net/docs.php",
    "[2] Chart.js Documentation, https://www.chartjs.org/docs/latest/",
    "[3] MDN Web Docs: Fetch API, https://developer.mozilla.org/",
    "[4] James et al., An Introduction to Statistical Learning, Springer.",
    "[5] Sommerville, Software Engineering, Pearson.",
]
for ref in refs:
    add_para(doc, ref)

output_path = "Expense_Tracker_IEEE_Report.docx"
doc.save(output_path)
print(f"Created: {output_path}")
